import os
from tkinter import messagebox
from docx import Document
from docx.shared import Inches
from utils import set_row_height_table, apply_table_styles, table_sort_center

def create_document(title_template, input_number, shuffled_list, images_output_folder):
  doc_question = Document()
  doc_empty = Document()
  doc_answer = Document()

  # 표 생성
  question_table = doc_question.add_table(rows=input_number + 2, cols=1)
  empty_table = doc_empty.add_table(rows=input_number + 2, cols=1)
  answer_table = doc_answer.add_table(rows=input_number + 2, cols=1)

  # 표 스타일
  question_table.style = doc_question.styles['Table Grid']
  empty_table.style = doc_empty.styles['Table Grid']
  answer_table.style = doc_answer.styles['Table Grid']

  set_row_height_table(question_table)
  set_row_height_table(empty_table)
  set_row_height_table(answer_table)

  # 표에 데이터 입력
  question_table_rows = question_table.rows
  empty_table_rows = empty_table.rows
  answer_table_rows = answer_table.rows

  for i in range(input_number + 2):
    question_data = question_table_rows[i].cells[0]
    empty_data = empty_table_rows[i].cells[0]
    answer_data = answer_table_rows[i].cells[0]

    if i == 0:
      question_data.text = title_template + '문제지'
      empty_data.text = title_template + '답안지'
      answer_data.text = title_template + '정답지'

      table_sort_center(doc_question, i)
      table_sort_center(doc_empty, i)
      table_sort_center(doc_answer, i)
      continue
    elif i == 1:
      name_text = '성명: '
      question_data.text = name_text
      empty_data.text = name_text
      answer_data.text = name_text
      continue

    output_number = i - 1
    question = shuffled_list[output_number - 1]
    output_number_template = f"<{output_number}번>"

    question_data.text = output_number_template

    # 문제에 이미지 포함 여부 확인
    if question['<제시그림>']:
      for extension in ['jpg', 'png']:
        image_path = f"{question['<제시그림>']}{extension}"

        if os.path.exists(image_path):
          run = question_data.add_paragraph().add_run()
          run.add_picture(image_path, width=Inches(3))  # 이미지 크기는 필요에 따라 조절
          break

    question_data.add_paragraph(f"{question['<문제>']}")

    if question['유형'] == '객관식':
      question_data.add_paragraph(f"\n{question['<답가지>']}")

    empty_data.text = output_number_template

    # 정답 이미지 or 텍스트
    answer_data.text = f"{output_number_template}"

    if question['<정답>'].startswith(images_output_folder):
      for extension in ['jpg', 'png']:
        image_path = f"{question['<정답>']}{extension}"

        if os.path.exists(image_path):
          run = answer_data.add_paragraph().add_run()
          run.add_picture(image_path, width=Inches(3))
          break
    else:
      answer_data.text += f"\n{question['<정답>']}"

  apply_table_styles(question_table)
  apply_table_styles(empty_table)
  apply_table_styles(answer_table)

  return doc_question, doc_empty, doc_answer

def save_documents(doc_question, doc_empty, doc_answer, input_title):
  try:
    question_path = f'[{input_title}] 문제지.docx'
    doc_question.save(question_path)

    empty_path = f'[{input_title}] 답안지.docx'
    doc_empty.save(empty_path)

    answer_path = f'[{input_title}] 정답지.docx'
    doc_answer.save(answer_path)

    messagebox.showinfo('성공', '모의고사 생성에 성공했습니다.')
  except PermissionError:
    messagebox.showerror('오류', '열려 있는 파일을 닫고, 다시 시도해주세요.')
    return
