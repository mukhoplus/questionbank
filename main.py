import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import random


def set_row_height_table(table):
  def set_row_height(row, height):
    tr = row._tr  # 행의 XML 요소
    trPr = tr.get_or_add_trPr()  # trPr 요소 추가
    trHeight = OxmlElement('w:trHeight')  # 높이 요소 생성
    trHeight.set(qn('w:val'), str(height))  # 높이 값 설정
    trHeight.set(qn('w:hRule'), 'exact')  # 높이 규칙 설정 ('atLeast' 또는 'exact' 가능)
    trPr.append(trHeight)  # trPr에 높이 요소 추가

  i = 0
  for row in table.rows:
    height = Pt(0.06) if i < 2 else Pt(0.09)
    set_row_height(row, height)
    i += 1

def set_font(cell, font_name, font_size, bold=False):
  for paragraph in cell.paragraphs:
    run = paragraph.runs[0]
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    paragraph.style.font.name = font_name
    paragraph.style.font.size = Pt(font_size)
    paragraph.style.font.bold = bold

def apply_table_styles(table):
  header_cells = table.rows[0].cells
  for cell in header_cells:
    set_font(cell, '굴림', 13, bold=True)

  for row in table.rows[1:]:
    for cell in row.cells:
        set_font(cell, '굴림', 11)

def get_answer_info(question_type, selections, answer):
  if question_type != '객관식':
    return answer
  selection_list = list(selections.split('\n'))
  
  for selection in selection_list:
    if answer in selection:
      return selection

def table_sort_center(document, row_index):
  document.tables[row_index].rows[0].cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

input_title = '테스트과목' # input('(테스트용)파일명을 입력하세요: ')
input_number = 5 # input('(테스트용)문제 수를 입력하세요: ')

# 엑셀 파일에서 데이터 불러오기
excel_path = f'./{input_title}.xlsx'  # 엑셀 파일 경로
sheet_name = '문제은행'  # 시트 이름

question_bank = []

# 기능 1: 파일 읽어서 파싱
try:
  df = pd.read_excel(excel_path, sheet_name=sheet_name)

  for index, row in df.iterrows():
    q_number = row['번호']
    q_type = row['유형']
    q_exam = row['시험']
    q_professor = row['교수님']
    q_year = row['<출제 년도>']
    q_question = row['<문제>']
    q_select = row['<답가지>']
    q_image = row['<제시그림>']
    q_answer = row['<정답>']
    q_description = row['<족보 페이지 또는 해설>']
    q_inspection = row['검수용\n번호']

    question = {
      '번호': q_number,
      '유형': q_type,
      '시험': q_exam,
      '교수님': q_professor,
      '<출제 년도>': q_year,
      '<문제>': q_question,
      '<답가지>': q_select if not pd.isna(q_select) else '',
      '<제시그림>': q_image if not pd.isna(q_image) else '',
      '<정답>': q_answer,
      '<족보 페이지 또는 해설>': q_description if not pd.isna(q_description)else '',
      '검수용 번호': q_inspection if not pd.isna(q_inspection) else ''
    }

    question_bank.append(question)
except FileNotFoundError:
  raise FileNotFoundError('해당 파일이 없습니다.')

# 기능 2: 문제지/답안지 만들기
limit = len(question_bank)
if limit == 0: raise ValueError('저장된 문제가 없습니다.')
if input_number > limit: input_number = limit

shuffled_list = random.sample(question_bank, limit)

# 제목 생성
year = 2024 # int(input('년도를 입력하세요: '))
semester = 2 # input('학기를 입력하세요: ')
exam_type = '중간' # input('시험 종류를 입력하세요(중간/기말): ')

title_template = f'{year}학년도 {semester}학기 {exam_type}고사 대비 모의고사 {input_title} 과목 '

# Word 문서 생성(문제지/답안지/정답지)
doc_question = Document()
doc_empty = Document()
doc_answer = Document()

# 표 생성
question_table = doc_question.add_table(rows = input_number + 2, cols = 1)
empty_table = doc_empty.add_table(rows = input_number + 2, cols = 1)
answer_table = doc_answer.add_table(rows = input_number + 2, cols = 1)

# 표 스타일
question_table.style = doc_question.styles['Table Grid']
empty_table.style =  doc_empty.styles['Table Grid']
answer_table.style = doc_answer.styles['Table Grid']

set_row_height_table(question_table)
set_row_height_table(empty_table)
set_row_height_table(answer_table)

# 표에 데이터 입력
question_table_rows = question_table.rows
empty_table_rows = empty_table.rows
answer_table_rows = answer_table.rows

for i in range(input_number + 2):
  question_data = question_table_rows[i].cells[0]
  empty_data = empty_table_rows[i].cells[0]
  answer_data = answer_table_rows[i].cells[0]

  if i == 0:
    question_data.text = title_template + '문제지'
    empty_data.text = title_template + '답안지'
    answer_data.text = title_template + '정답지'

    table_sort_center(doc_question, i)
    table_sort_center(doc_empty, i)
    table_sort_center(doc_answer, i)
    continue
  elif i == 1:
    name_text = '성명: '
    question_data.text = name_text
    empty_data.text = name_text
    answer_data.text = name_text
    continue
  
  output_number = i - 1 # 만들어진 문제의 순서
  question, table_index = shuffled_list[output_number - 1], i
  
  output_number_template = f"<{output_number}번>"

  question_data.text = f"{output_number_template}\n{question['<문제>']}"
  empty_data.text = output_number_template
  answer_data.text = f"{output_number_template}\n{get_answer_info(question['유형'], question['<답가지>'], question['<정답>'])}"

# 글꼴 설정
apply_table_styles(question_table)
apply_table_styles(empty_table)
apply_table_styles(answer_table)

# 파일 저장
try:
  question_path = f'[{input_title}] 문제지.docx'
  doc_question.save(question_path)
  print(f"문제지 파일이 {question_path}에 저장되었습니다.")

  empty_path = f'[{input_title}] 답안지.docx'
  doc_empty.save(empty_path)
  print(f"답안지 파일이 {empty_path}에 저장되었습니다.")

  answer_path = f'[{input_title}] 정답지.docx'
  doc_answer.save(answer_path)
  print(f"정답지 파일이 {answer_path}에 저장되었습니다.")
except PermissionError:
  print('열려 있는 파일을 닫고, 다시 시도해주세요.')
