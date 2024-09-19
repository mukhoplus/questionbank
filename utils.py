from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from tkinter import messagebox

def set_row_height_table(table):
  def set_row_height(row, min_height):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(min_height))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

  i = 0
  for row in table.rows:
    min_height = Pt(0.06) if i < 2 else Pt(0.09)
    set_row_height(row, min_height)
    i += 1

def set_font(cell, font_name, font_size, bold=False):
  for paragraph in cell.paragraphs:
    try:
      run = paragraph.runs[0]
      run.font.name = font_name
      run.font.size = Pt(font_size)
      run.font.bold = bold
      paragraph.style.font.name = font_name
      paragraph.style.font.size = Pt(font_size)
      paragraph.style.font.bold = bold
    except:
      pass

def apply_table_styles(table):
  header_cells = table.rows[0].cells
  for cell in header_cells:
    set_font(cell, 'NS regular', 13, bold=True)

  for row in table.rows[1:]:
    for cell in row.cells:
      set_font(cell, 'NS regular', 11)

def table_sort_center(document, row_index):
  document.tables[row_index].rows[0].cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

def handling_input_exception(input_title, input_number, year):
  if not input_title:
    messagebox.showwarning('경고', '과목명을 입력해주세요.')
    return False
  
  if not input_number:
    messagebox.showwarning('경고', '문제 수를 입력해주세요.')
    return False
  
  if not year:
    messagebox.showwarning('경고', '학년도를 입력해주세요.')
    return False
  
  try:
    error_check = int(input_number)
    if type(error_check) != int:
      raise ValueError
    if int(input_number) < 1:
      raise ValueError
  except:
    messagebox.showwarning('경고', '문제 수는 1 이상의 정수여야 합니다.')
    return False
  
  try:
    error_check = int(year)
    if type(error_check) != int:
      raise ValueError
    if error_check < 0:
      raise ValueError
  except:
    messagebox.showwarning('경고', '학년도는 0 이상의 정수여야 합니다.')
    return False

  return True

def get_subject(excel_path):
  path_list = excel_path.split('_')
  file_list = path_list[-1].split('.')
  subject = file_list[0]

  return subject
